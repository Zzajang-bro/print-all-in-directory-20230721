from os  import walk
from os.path  import join
from sys import argv

# pip install python-docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

if len(argv) < 3:
	tar_path = input('타겟 경로: ')
	proj_name = input('프로젝트 이름: ')
else:
	tar_path = argv[1]
	projName = argv[2]

doc = Document()

doc.add_heading(proj_name, level=0)

for (root, dirs, files) in walk(tar_path):
	for file in files:
		doc.add_heading(join(proj_name, root[len(tar_path):],file), level=1)
		with open(f'{join(root,file)}','r', encoding='utf-8', errors='ignore') as d:
			dat = d.read()
			doc.add_paragraph(dat)

doc.save('paid.result.docx')
